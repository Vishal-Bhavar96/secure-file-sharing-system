import io
import base64
import html
import mimetypes
from typing import Dict, Any, Optional

def extract_docx_html(raw_bytes: bytes) -> str:
    """Extract paragraphs, headings, bullet lists, and tables from a .docx file into sanitized HTML"""
    try:
        import docx
        doc = docx.Document(io.BytesIO(raw_bytes))
        html_parts = []

        # Process paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            safe_text = html.escape(text)
            style_name = (p.style.name if p.style else "").lower()
            if "heading 1" in style_name or "title" in style_name:
                html_parts.append(f"<h2 class='docx-heading-1'>{safe_text}</h2>")
            elif "heading 2" in style_name:
                html_parts.append(f"<h3 class='docx-heading-2'>{safe_text}</h3>")
            elif "heading 3" in style_name:
                html_parts.append(f"<h4 class='docx-heading-3'>{safe_text}</h4>")
            elif "list" in style_name or "bullet" in style_name:
                html_parts.append(f"<li class='docx-list-item'>{safe_text}</li>")
            else:
                html_parts.append(f"<p class='docx-paragraph'>{safe_text}</p>")

        # Process tables
        for table in doc.tables:
            table_html = ["<table class='docx-table'>"]
            for r_idx, row in enumerate(table.rows):
                table_html.append("<tr>")
                for cell in row.cells:
                    tag = "th" if r_idx == 0 else "td"
                    cell_text = html.escape(cell.text.strip())
                    table_html.append(f"<{tag}>{cell_text}</{tag}>")
                table_html.append("</tr>")
            table_html.append("</table>")
            html_parts.append("".join(table_html))

        if not html_parts:
            return "<div class='docx-empty-notice'><p><em>(The Word document contains no readable text or is empty)</em></p></div>"

        return "".join(html_parts)
    except Exception as e:
        return f"<div class='docx-error-notice'><p><em>Unable to render Word preview: {html.escape(str(e))}</em></p></div>"

def generate_file_preview(decrypted_bytes: bytes, filename: str, mime_type: Optional[str] = None) -> Dict[str, Any]:
    """Generates an in-browser online viewing payload for different file types"""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    size_bytes = len(decrypted_bytes)
    
    if not mime_type or mime_type == "application/octet-stream":
        guessed, _ = mimetypes.guess_type(filename)
        mime_type = guessed or "application/octet-stream"

    # 1. Images
    if ext in ("png", "jpg", "jpeg", "gif", "webp", "svg", "bmp", "ico") or mime_type.startswith("image/"):
        b64_data = base64.b64encode(decrypted_bytes).decode("ascii")
        data_uri = f"data:{mime_type};base64,{b64_data}"
        return {
            "preview_type": "image",
            "filename": filename,
            "mime_type": mime_type,
            "size_bytes": size_bytes,
            "data_uri": data_uri,
            "can_render": True
        }

    # 2. PDF Documents
    if ext == "pdf" or mime_type == "application/pdf":
        b64_data = base64.b64encode(decrypted_bytes).decode("ascii")
        data_uri = f"data:application/pdf;base64,{b64_data}"
        return {
            "preview_type": "pdf",
            "filename": filename,
            "mime_type": "application/pdf",
            "size_bytes": size_bytes,
            "data_uri": data_uri,
            "can_render": True
        }

    # 3. Word Documents (.docx)
    if ext == "docx" or "wordprocessingml" in mime_type:
        rendered_html = extract_docx_html(decrypted_bytes)
        return {
            "preview_type": "docx",
            "filename": filename,
            "mime_type": mime_type,
            "size_bytes": size_bytes,
            "html_content": rendered_html,
            "can_render": True
        }

    # 4. Text / Source Code / Structured Data
    text_extensions = {
        "txt", "md", "csv", "json", "xml", "yaml", "yml", "py", "js", "ts",
        "html", "css", "sql", "sh", "bat", "ps1", "env", "ini", "log", "cfg"
    }
    if ext in text_extensions or mime_type.startswith("text/") or mime_type in ("application/json", "application/xml"):
        try:
            text_content = decrypted_bytes.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text_content = decrypted_bytes.decode("latin-1")
            except Exception:
                text_content = None

        if text_content is not None:
            # Cap preview text at 500 KB for responsive client rendering
            is_truncated = len(text_content) > 500000
            display_text = text_content[:500000]
            lines = display_text.splitlines()
            return {
                "preview_type": "text",
                "filename": filename,
                "mime_type": mime_type,
                "size_bytes": size_bytes,
                "text_content": display_text,
                "line_count": len(lines),
                "is_truncated": is_truncated,
                "can_render": True
            }

    # 5. Media (Audio / Video)
    if ext in ("mp4", "webm", "ogg", "mp3", "wav") or mime_type.startswith(("video/", "audio/")):
        b64_data = base64.b64encode(decrypted_bytes).decode("ascii")
        data_uri = f"data:{mime_type};base64,{b64_data}"
        media_kind = "video" if (ext in ("mp4", "webm", "ogg") or mime_type.startswith("video/")) else "audio"
        return {
            "preview_type": "media",
            "media_kind": media_kind,
            "filename": filename,
            "mime_type": mime_type,
            "size_bytes": size_bytes,
            "data_uri": data_uri,
            "can_render": True
        }

    # 6. Other Binary / Unsupported
    return {
        "preview_type": "binary",
        "filename": filename,
        "mime_type": mime_type,
        "size_bytes": size_bytes,
        "can_render": False,
        "message": f"This file format ({ext.upper() or 'Binary'}) does not support in-browser visual preview."
    }
